"""Regression tests for v1.2 quality gates and cross-agent delivery."""
from __future__ import annotations
import json, sys, tempfile, unittest
from pathlib import Path
from docx import Document
ROOT=Path(__file__).resolve().parents[1]; sys.path.insert(0,str(ROOT/"scripts"))
from build_report import ReportBuildError, build_report
from qa_report import check_report, render_report
from validate_plan import PlanValidationError, validate_generation_plan
from dashboard_server import create_app

class V12QualityTests(unittest.TestCase):
 def test_numeric_score_requires_basis(self)->None:
  plan={"report_kind":"revision","source_state":"completed","summary":"综合 79/100","verdict":"修改后提交","additions":[]}
  with self.assertRaises(PlanValidationError): validate_generation_plan(plan)

 def test_strict_anchor_rejects_unlocated_content(self)->None:
  with tempfile.TemporaryDirectory() as temp:
   root=Path(temp); source=root/"source.docx"; Document().save(source)
   plan=root/"plan.json"; plan.write_text(json.dumps({"report_kind":"revision","source_state":"completed","summary":"s","verdict":"v","additions":[{"anchor_text":"不存在","category":"issue","text":"缺口"}]},ensure_ascii=False),encoding="utf-8")
   with self.assertRaises(ReportBuildError): build_report(source,plan,root/"out")

 def test_structured_blocks_single_docx_qa_and_download(self)->None:
  with tempfile.TemporaryDirectory() as temp:
   root=Path(temp); source=root/"source.docx"; doc=Document(); doc.add_paragraph("实验步骤"); doc.add_paragraph("实验结果"); doc.save(source)
   plan_data={"report_kind":"execution","source_state":"blank","summary":"按真实步骤执行。","verdict":"尚未完成","generate_companion":True,"additions":[
    {"anchor_text":"实验步骤","category":"guidance","priority":"high","evidence_basis":"template","block_type":"checklist","label":"执行","items":["确认环境","执行并记录"]},
    {"anchor_text":"实验结果","category":"evidence","priority":"medium","evidence_basis":"template","block_type":"table","label":"证据表","columns":["任务","证据"],"rows":[["配置","截图"]]}
   ]}
   plan=root/"plan.json"; plan.write_text(json.dumps(plan_data,ensure_ascii=False),encoding="utf-8")
   output,metadata_path=build_report(source,plan,root/"out"); generated=Document(output); text="\n".join(p.text for p in generated.paragraphs)
   self.assertIn("☐ 确认环境",text); self.assertNotIn("未定位内容",text); self.assertEqual(len(generated.tables),2)
   metadata=json.loads(metadata_path.read_text(encoding="utf-8")); self.assertEqual(len(metadata["generated_files"]),1); self.assertEqual(metadata["generated_files"][0]["kind"],"annotated")
   quality,quality_path=check_report(source,output,metadata_path); self.assertTrue(quality["ok"]); self.assertTrue(quality_path.is_file()); self.assertIn("自动质量检查","\n".join(p.text for p in Document(output).paragraphs))
   metadata=json.loads(metadata_path.read_text(encoding="utf-8")); metadata["quality"]=quality; metadata_path.write_text(json.dumps(metadata,ensure_ascii=False),encoding="utf-8")
   client=create_app(root/"out").test_client(); listing=client.get("/api/reports").get_json()["reports"][0]; self.assertEqual(len(listing["files"]),1)
   download=client.get(listing["files"][0]["download_url"]); self.assertEqual(download.status_code,200); download.close()
 def test_render_unavailable_is_explicit(self)->None:
  from unittest.mock import patch
  with tempfile.TemporaryDirectory() as temp:
   doc=Path(temp)/"report.docx"; Document().save(doc)
   with patch("qa_report._render_external",return_value=(False,"missing")), patch("qa_report._render_soffice",return_value=(False,"missing")), patch.dict("os.environ",{},clear=True):
    result=render_report(doc,"job")
   self.assertEqual(result["status"],"unavailable"); self.assertEqual(len(result["attempts"]),3)
   self.assertIn("Skipped by default",result["attempts"][2]["message"])
 def test_openclaw_contract_is_present(self)->None:
  skill=(ROOT/"SKILL.md").read_text(encoding="utf-8"); metadata_line=next(line for line in skill.splitlines() if line.startswith("metadata: "))
  json.loads(metadata_line.removeprefix("metadata: ")); self.assertIn("OpenClaw",skill)
  install_sh=ROOT/"install.sh"; install_ps1=ROOT/"install.ps1"
  if install_sh.exists(): self.assertIn(".openclaw/skills",install_sh.read_text(encoding="utf-8"))
  if install_ps1.exists(): self.assertIn(".openclaw\\skills",install_ps1.read_text(encoding="utf-8"))

if __name__=="__main__": unittest.main()
