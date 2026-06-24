from __future__ import annotations
import json
import sys
import tempfile
import unittest
from pathlib import Path
from docx import Document
from PIL import Image

ROOT=Path(__file__).resolve().parents[1]
sys.path.insert(0,str(ROOT/"scripts"))
from build_report import build_report
from dashboard_server import create_app
from inspect_report import prepare_report
from review_signals import classify_report_signals
from validate_plan import PlanValidationError, validate_generation_plan

class ReviewSignalTests(unittest.TestCase):
 def test_blank_partial_false_complete_and_completed(self)->None:
  self.assertEqual(classify_report_signals("请填写实验步骤和实验结果",0)["suggested_state"],"blank")
  self.assertEqual(classify_report_signals("实验步骤：配置软件并启动任务。",0)["suggested_state"],"partial")
  false_complete=classify_report_signals("实验步骤：执行任务。实验结论：任务已经完成。",0)
  self.assertEqual(false_complete["suggested_state"],"partial")
  self.assertTrue(false_complete["false_completion_findings"])
  completed=classify_report_signals("实验步骤：执行任务。实验结果：输出数据 42。实验结论：结果符合目标。",1)
  self.assertEqual(completed["suggested_state"],"completed")

 def test_contamination_and_small_screenshot_hint(self)->None:
  signals=classify_report_signals("以下是 ChatGPT 生成的报告，请填写实验结果",0)
  self.assertGreaterEqual(len(signals["contamination_findings"]),2)
  with tempfile.TemporaryDirectory() as temp:
   root=Path(temp); image=root/"small.png"; Image.new("RGB",(320,180),"white").save(image)
   manifest=json.loads(prepare_report(image,root/"prepared").read_text(encoding="utf-8"))
   self.assertEqual(manifest["visuals"][0]["review_hint"],"retake-or-original-resolution")

class WorkbenchContractTests(unittest.TestCase):
 def test_budget_limit_is_enforced(self)->None:
  plan={"report_kind":"revision","source_state":"partial","submission_signal":"yellow","time_budget":"15m","estimated_minutes":20,"summary":"s","verdict":"v","additions":[]}
  with self.assertRaises(PlanValidationError): validate_generation_plan(plan)

 def test_guidance_only_does_not_generate_docx(self)->None:
  with tempfile.TemporaryDirectory() as temp:
   root=Path(temp); source=root/"source.docx"; Document().save(source); plan=root/"plan.json"
   plan.write_text(json.dumps({"report_kind":"execution","source_state":"blank","output_mode":"guidance_only","summary":"s","verdict":"v","additions":[]}),encoding="utf-8")
   from build_report import ReportBuildError
   with self.assertRaises(ReportBuildError): build_report(source,plan,root/"out")

 def test_single_docx_budget_actions_and_safe_screenshot_route(self)->None:
  with tempfile.TemporaryDirectory() as temp:
   root=Path(temp); source=root/"source.docx"; document=Document()
   for index in range(6): document.add_paragraph(f"位置 {index}")
   document.save(source)
   screenshot=root/"evidence.png"; Image.new("RGB",(1000,700),"white").save(screenshot)
   additions=[]
   priorities=["low","blocker","high","medium","high","optional"]
   for index,priority in enumerate(priorities):
    additions.append({"anchor_text":f"位置 {index}","category":"suggestion","priority":priority,"estimated_minutes":5,"evidence_basis":"text","text":f"动作 {index}"})
   plan={"report_kind":"revision","source_state":"partial","submission_signal":"yellow","time_budget":"15m","estimated_minutes":15,"summary":"需要补证据","verdict":"小修后提交","additions":additions,
    "false_completion_findings":[{"location":"结果","issue":"证据不足","evidence":"无运行输出","action":"补运行结果","severity":"high"}],
    "contamination_findings":[{"location":"正文","issue":"占位符","evidence":"发现待补充","action":"删除占位符","severity":"medium"}],
    "screenshot_evidence":[{"source_ref":"图1","claim":"证明运行成功","observed":"仅看到主界面","readability":"partial","matches_text":"no","privacy_risk":"无","retake_instruction":"打开结果窗口并包含状态、时间和关键字段","image_path":str(screenshot),"embed_in_appendix":False}]}
   plan_path=root/"plan.json"; plan_path.write_text(json.dumps(plan,ensure_ascii=False),encoding="utf-8")
   output,metadata_path=build_report(source,plan_path,root/"out"); metadata=json.loads(metadata_path.read_text(encoding="utf-8"))
   self.assertEqual(len(metadata["generated_files"]),1)
   self.assertLessEqual(len(metadata["actions"]),5); self.assertEqual(metadata["estimated_minutes"],15)
   self.assertEqual([item["priority"] for item in metadata["actions"]],["blocker","high","high"])
   self.assertIn("Codex 综合审阅附录","\n".join(p.text for p in Document(output).paragraphs))
   client=create_app(root/"out").test_client(); listing=client.get("/api/reports").get_json()["reports"][0]
   self.assertEqual(listing["submission_signal"],"yellow"); self.assertEqual(len(listing["files"]),1)
   response=client.get(listing["screenshots"][0]["preview_url"]); self.assertEqual(response.status_code,200); response.close()
   metadata["screenshot_evidence"][0]["preview_name"]="../evidence.png"; metadata_path.write_text(json.dumps(metadata,ensure_ascii=False),encoding="utf-8")
   self.assertEqual(client.get(f"/api/reports/{metadata['job_id']}/screenshots/0").status_code,404)

if __name__=="__main__": unittest.main()
