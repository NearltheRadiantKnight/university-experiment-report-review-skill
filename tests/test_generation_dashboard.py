"""Tests for generated DOCX styling and local dashboard delivery."""

from __future__ import annotations

import hashlib
import json
import sys
import tempfile
import unittest
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor

SCRIPTS_DIR = Path(__file__).resolve().parents[1] / "scripts"
sys.path.insert(0, str(SCRIPTS_DIR))

from build_report import build_report
from dashboard_server import DASHBOARD_API_VERSION, DASHBOARD_ASSET_VERSION, create_app
from dashboard_launcher import API_VERSION, ASSET_VERSION


def file_hash(path: Path) -> str:
    """Return a stable SHA-256 digest for immutability assertions."""
    return hashlib.sha256(path.read_bytes()).hexdigest()


def create_source(path: Path) -> None:
    """Create a styled source DOCX with stable review anchors."""
    document = Document()
    paragraph = document.add_paragraph()
    run = paragraph.add_run("实验步骤")
    run.font.name = "Arial"
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x11, 0x22, 0x33)
    run.bold = True
    document.add_paragraph("实验结果")
    document.add_paragraph("实验结论")
    document.save(path)


class GeneratedReportTests(unittest.TestCase):
    def test_dashboard_launcher_api_matches_server(self) -> None:
        self.assertEqual(API_VERSION, DASHBOARD_API_VERSION)
        self.assertEqual(ASSET_VERSION, DASHBOARD_ASSET_VERSION)

    """Verify the original remains unchanged and additions are differentiated."""

    def test_execution_report_preserves_original_run(self) -> None:
        """Original font settings must survive while guidance is blue and labeled."""
        with tempfile.TemporaryDirectory() as temp_name:
            root = Path(temp_name)
            source = root / "blank.docx"
            plan = root / "plan.json"
            output_dir = root / "outputs"
            create_source(source)
            source_before = file_hash(source)
            plan.write_text(
                json.dumps(
                    {
                        "report_kind": "execution",
                        "source_state": "blank",
                        "summary": "按真实步骤完成实验。",
                        "verdict": "尚未完成",
                        "additions": [
                            {
                                "anchor_text": "实验步骤",
                                "category": "guidance",
                                "label": "执行顺序",
                                "text": "先确认环境，再执行并记录真实结果。",
                            }
                        ],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )

            output_path, metadata_path = build_report(source, plan, output_dir)
            generated = Document(output_path)
            original_run = generated.paragraphs[0].runs[0]
            inserted = generated.paragraphs[1]

            self.assertEqual(file_hash(source), source_before)
            self.assertEqual(original_run.font.name, "Arial")
            self.assertEqual(original_run.font.size.pt, 13)
            self.assertEqual(str(original_run.font.color.rgb), "112233")
            self.assertTrue(original_run.bold)
            self.assertIn("Codex 新增", inserted.text)
            self.assertEqual(str(inserted.runs[0].font.color.rgb), "2F75B5")
            self.assertEqual(inserted.runs[0].font.name, "Microsoft YaHei")
            metadata = json.loads(metadata_path.read_text(encoding="utf-8"))
            self.assertEqual(metadata["report_kind"], "execution")
            self.assertIn("实验执行报告", output_path.name)

    def test_revision_order_and_dashboard_download(self) -> None:
        """Revision notes should stay ordered and be downloadable from loopback app."""
        with tempfile.TemporaryDirectory() as temp_name:
            root = Path(temp_name)
            source = root / "completed.docx"
            plan = root / "plan.json"
            output_dir = root / "outputs"
            create_source(source)
            plan.write_text(
                json.dumps(
                    {
                        "report_kind": "revision",
                        "source_state": "completed",
                        "summary": "截图证据需要补强。",
                        "verdict": "小修后可提交",
                        "additions": [
                            {"anchor_text": "实验结果", "category": "issue", "text": "关键参数不可读。"},
                            {"anchor_text": "实验结果", "category": "suggestion", "text": "替换为原始分辨率截图。"},
                        ],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )

            output_path, metadata_path = build_report(source, plan, output_dir)
            generated = Document(output_path)
            texts = [paragraph.text for paragraph in generated.paragraphs]
            result_index = texts.index("实验结果")
            self.assertIn("发现问题", texts[result_index + 1])
            self.assertIn("修改建议", texts[result_index + 2])
            self.assertIn("修改报告", output_path.name)

            app = create_app(output_dir)
            app.testing = True
            client = app.test_client()
            page = client.get("/")
            self.assertEqual(page.status_code, 200)
            self.assertIn("实验报告生成中心".encode("utf-8"), page.data)
            health = client.get("/api/health").get_json()
            self.assertEqual(health["service"], "university-experiment-report-dashboard")
            self.assertEqual(health["api_version"], 4)
            self.assertEqual(health["asset_version"], "1.5.7")
            self.assertTrue(health["output_dir_id"])
            listing = client.get("/api/reports")
            self.assertEqual(listing.status_code, 200)
            payload = listing.get_json()
            self.assertEqual(len(payload["reports"]), 1)
            job_id = json.loads(metadata_path.read_text(encoding="utf-8"))["job_id"]
            download = client.get(f"/api/reports/{job_id}/download")
            self.assertEqual(download.status_code, 200)
            self.assertGreater(len(download.data), 100)
            download.close()
            self.assertEqual(client.get("/api/reports/not-valid/download").status_code, 404)


class DashboardFrontendTests(unittest.TestCase):
    def test_progressive_workbench_avoids_background_refresh_and_html_injection(self) -> None:
        script = (SCRIPTS_DIR.parent / "assets" / "dashboard" / "static" / "app.js").read_text(encoding="utf-8-sig")
        self.assertIn("report.download_url", script)
        self.assertIn("feedbackDirty", script)
        self.assertIn("window.confirm", script)
        self.assertIn("loadReports(false)", script)
        self.assertNotIn("setInterval", script)
        self.assertNotIn("innerHTML", script)
        self.assertIn("textContent", script)
        self.assertIn("openGenerationSettings", script)
        self.assertNotIn("queueSkillImprovement", script)
        self.assertNotIn("skillImproveBtn", script)
        self.assertIn("delete-feedback", script)
        self.assertIn("clear-feedback", script)
        self.assertNotIn("feedback-status", script)
        self.assertIn("retry-render", script)
        self.assertIn("dirtyEditors", script)
        self.assertIn("markSaved", script)
        self.assertIn("personal-memory", script)
        self.assertIn('screenshotMatchLabel(item.matches_text)', script)
        self.assertIn("\u4e2a\u6027\u5316\u8bb0\u5fc6", script)
        self.assertIn("\\u5f85\\u786e\\u8ba4", script)
        self.assertNotIn("\u4e0b\u8f7d\u53cd\u9988 JSON", script)
        self.assertNotIn("feedback-download", script)
        self.assertIn('remove.type="button"', script)
        self.assertIn('close.type="button"', script)
        self.assertIn('aria-label","\u5173\u95ed\u5f39\u7a97"', script)
        self.assertIn('role","dialog"', script)
        self.assertIn('aria-modal","true"', script)
        self.assertIn('showRender=status==="passed"||status==="failed"||status==="permission-required"||hasPreview', script)
        self.assertNotIn("\u6e32\u67d3\u5668\u4e0d\u53ef\u7528", script)
        template = (SCRIPTS_DIR.parent / "assets" / "dashboard" / "templates" / "index.html").read_text(encoding="utf-8")
        self.assertIn("v='1.5.7'", template)
        self.assertNotIn("skillImproveBtn", template)

    def test_dashboard_button_css_contract(self) -> None:
        styles = (SCRIPTS_DIR.parent / "assets" / "dashboard" / "static" / "styles.css").read_text(encoding="utf-8")
        self.assertIn("inline-size: 36px", styles)
        self.assertIn(".modal-overlay[hidden]", styles)
        self.assertIn("display: none !important", styles)
        self.assertIn("min-inline-size:64px", styles)
        self.assertIn("white-space:nowrap", styles)
        self.assertNotIn("width: 32px; height: 32px", styles)
        self.assertIn("grid-template-columns:minmax(180px,1.1fr) minmax(112px,132px) minmax(280px,2fr) auto", styles)
        self.assertNotIn(".render-state.unavailable", styles)
        self.assertNotIn("background:#fff6df", styles)

if __name__ == "__main__":
    unittest.main()
