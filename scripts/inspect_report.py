#!/usr/bin/env python3
"""Prepare a university experiment report for local agent review.

The script extracts text and visual evidence from DOCX, PDF, plain-text,
Markdown, and image files without calling any network service or model API.
It writes a manifest that the agent can read before performing semantic review.
"""

from __future__ import annotations

import argparse
import hashlib
import json
import shutil
import sys
import zipfile
from pathlib import Path
from typing import Any
from review_signals import classify_report_signals
from xml.etree import ElementTree

SKILL_NAME = "university-experiment-report-review-skill"
VERSION = "1.5.0"
SUPPORTED_SUFFIXES = {".docx", ".pdf", ".txt", ".md", ".png", ".jpg", ".jpeg", ".bmp", ".gif", ".webp", ".tif", ".tiff"}
IMAGE_SUFFIXES = {".png", ".jpg", ".jpeg", ".bmp", ".gif", ".webp", ".tif", ".tiff"}


class ReportPreparationError(RuntimeError):
    """Raised when a report cannot be prepared safely for local review."""


def _structured_error(message: str, error_type: str, hint: str) -> None:
    payload = {"error": message, "error_type": error_type, "hint": hint}
    print(json.dumps(payload, ensure_ascii=False), file=sys.stderr)


def _sha256(file_path: Path) -> str:
    digest = hashlib.sha256()
    with file_path.open("rb") as input_file:
        for block in iter(lambda: input_file.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def _safe_name(name: str, fallback: str) -> str:
    cleaned = "".join(character if character.isalnum() or character in ".-_" else "_" for character in name)
    return cleaned or fallback


def _write_text(text: str, output_dir: Path) -> Path:
    text_path = output_dir / "document.txt"
    text_path.write_text(text, encoding="utf-8")
    return text_path


def _docx_text(xml_bytes: bytes) -> str:
    try:
        root = ElementTree.fromstring(xml_bytes)
    except ElementTree.ParseError as exc:
        raise ReportPreparationError(f"DOCX XML could not be parsed: {exc}") from exc

    paragraphs: list[str] = []
    for paragraph in root.iter("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p"):
        fragments = [
            node.text or ""
            for node in paragraph.iter("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t")
        ]
        paragraph_text = "".join(fragments).strip()
        if paragraph_text:
            paragraphs.append(paragraph_text)
    return "\n".join(paragraphs)


def _docx_visual_contexts(input_path: Path) -> dict[str, list[dict[str, Any]]]:
    """Map embedded media to nearby source paragraphs using DOCX relationships."""
    try:
        from docx import Document
        from docx.oxml.ns import qn
    except ImportError:
        return {}
    try:
        with zipfile.ZipFile(input_path) as archive:
            if "[Content_Types].xml" not in archive.namelist():
                return {}
        document = Document(str(input_path))
    except Exception:
        return {}
    contexts: dict[str, list[dict[str, Any]]] = {}
    for index, paragraph in enumerate(document.paragraphs):
        nearby = " | ".join(
            item.text.strip()
            for item in document.paragraphs[max(0, index - 1): min(len(document.paragraphs), index + 2)]
            if item.text.strip()
        )
        for blip in paragraph._p.xpath(".//a:blip"):
            relationship_id = blip.get(qn("r:embed"))
            if not relationship_id or relationship_id not in document.part.related_parts:
                continue
            part_name = str(document.part.related_parts[relationship_id].partname).lstrip("/")
            contexts.setdefault(part_name, []).append({"paragraph_index": index, "nearby_text": nearby[:500]})
    return contexts


def _create_contact_sheet(output_dir: Path, visuals: list[dict[str, Any]], warnings: list[str]) -> str | None:
    """Create a local thumbnail overview without OCR or network access."""
    if not visuals:
        return None
    try:
        from PIL import Image, ImageDraw, ImageFont
    except ImportError:
        warnings.append("Pillow is unavailable; contact sheet was not created.")
        return None
    tile_width, tile_height, columns = 520, 390, 2
    rows = (len(visuals) + columns - 1) // columns
    sheet = Image.new("RGB", (tile_width * columns, tile_height * rows), "white")
    draw = ImageDraw.Draw(sheet)
    font = ImageFont.load_default()
    for index, visual in enumerate(visuals):
        image_path = output_dir / str(visual["path"])
        try:
            with Image.open(image_path) as image:
                image = image.convert("RGB")
                image.thumbnail((tile_width - 30, tile_height - 55))
                x = (index % columns) * tile_width + (tile_width - image.width) // 2
                y = (index // columns) * tile_height + 28
                sheet.paste(image, (x, y))
        except Exception as exc:
            warnings.append(f"Could not add {image_path.name} to contact sheet: {exc}")
        label = f"{index + 1}. {Path(str(visual['path'])).name}"
        draw.text(((index % columns) * tile_width + 12, (index // columns) * tile_height + 8), label, fill="#17324d", font=font)
    overview = output_dir / "contact-sheet.jpg"
    sheet.save(overview, quality=88)
    return overview.relative_to(output_dir).as_posix()

def _prepare_docx(input_path: Path, output_dir: Path) -> tuple[str, list[dict[str, Any]], list[str]]:
    images_dir = output_dir / "images"
    images_dir.mkdir(parents=True, exist_ok=True)
    text_sections: list[str] = []
    visuals: list[dict[str, Any]] = []
    warnings: list[str] = []

    try:
        with zipfile.ZipFile(input_path) as archive:
            names = archive.namelist()
            xml_names = [
                name
                for name in names
                if name == "word/document.xml"
                or name.startswith("word/header") and name.endswith(".xml")
                or name.startswith("word/footer") and name.endswith(".xml")
                or name in {"word/footnotes.xml", "word/endnotes.xml"}
            ]
            for xml_name in xml_names:
                extracted = _docx_text(archive.read(xml_name))
                if extracted:
                    text_sections.append(f"## {xml_name}\n{extracted}")

            media_names = [name for name in names if name.startswith("word/media/")]
            for index, media_name in enumerate(sorted(media_names), start=1):
                suffix = Path(media_name).suffix.lower()
                if suffix not in IMAGE_SUFFIXES:
                    warnings.append(f"Skipped unsupported embedded media: {media_name}")
                    continue
                output_name = f"docx-image-{index:03d}-{_safe_name(Path(media_name).name, f'image{suffix}')}"
                image_path = images_dir / output_name
                image_path.write_bytes(archive.read(media_name))
                visuals.append(
                    {
                        "kind": "embedded-image",
                        "source": media_name,
                        "path": image_path.relative_to(output_dir).as_posix(),
                    }
                )
    except zipfile.BadZipFile as exc:
        raise ReportPreparationError("The DOCX file is damaged or is not a valid DOCX archive.") from exc

    contexts = _docx_visual_contexts(input_path)
    for visual in visuals:
        visual["contexts"] = contexts.get(str(visual.get("source", "")), [])
    if not text_sections:
        warnings.append("No extractable DOCX text was found; inspect the exported images and original document visually.")
    return "\n\n".join(text_sections), visuals, warnings


def _prepare_pdf(
    input_path: Path,
    output_dir: Path,
    max_pages: int,
) -> tuple[str, list[dict[str, Any]], list[str]]:
    try:
        import fitz
    except ImportError as exc:
        raise ReportPreparationError(
            "PDF preparation requires PyMuPDF. Use your agent's native PDF viewer or install the declared dependency."
        ) from exc

    images_dir = output_dir / "images"
    images_dir.mkdir(parents=True, exist_ok=True)
    text_sections: list[str] = []
    visuals: list[dict[str, Any]] = []
    warnings: list[str] = []
    seen_xrefs: set[int] = set()

    try:
        document = fitz.open(input_path)
    except Exception as exc:
        raise ReportPreparationError(f"PDF could not be opened: {exc}") from exc

    page_limit = min(len(document), max_pages)
    if len(document) > max_pages:
        warnings.append(f"PDF has {len(document)} pages; only the first {max_pages} were prepared.")

    try:
        for page_index in range(page_limit):
            page = document.load_page(page_index)
            page_number = page_index + 1
            page_text = page.get_text("text").strip()
            if page_text:
                text_sections.append(f"## Page {page_number}\n{page_text}")

            rendered_path = images_dir / f"page-{page_number:03d}.png"
            page.get_pixmap(matrix=fitz.Matrix(1.5, 1.5), alpha=False).save(rendered_path)
            visuals.append(
                {
                    "kind": "page-render",
                    "page": page_number,
                    "path": rendered_path.relative_to(output_dir).as_posix(),
                }
            )

            for image_index, image_info in enumerate(page.get_images(full=True), start=1):
                xref = int(image_info[0])
                if xref in seen_xrefs:
                    continue
                seen_xrefs.add(xref)
                image_data = document.extract_image(xref)
                extension = _safe_name(str(image_data.get("ext", "png")), "png")
                embedded_path = images_dir / f"page-{page_number:03d}-image-{image_index:02d}.{extension}"
                embedded_path.write_bytes(image_data["image"])
                visuals.append(
                    {
                        "kind": "embedded-image",
                        "page": page_number,
                        "xref": xref,
                        "path": embedded_path.relative_to(output_dir).as_posix(),
                    }
                )
    finally:
        document.close()

    if not text_sections:
        warnings.append("No PDF text was extracted; treat it as a scanned document and inspect the page renders.")
    return "\n\n".join(text_sections), visuals, warnings


def _prepare_text(input_path: Path) -> tuple[str, list[dict[str, Any]], list[str]]:
    try:
        return input_path.read_text(encoding="utf-8"), [], []
    except UnicodeDecodeError:
        return input_path.read_text(encoding="utf-8", errors="replace"), [], [
            "Invalid UTF-8 bytes were replaced while reading the text file."
        ]


def _prepare_image(input_path: Path, output_dir: Path) -> tuple[str, list[dict[str, Any]], list[str]]:
    images_dir = output_dir / "images"
    images_dir.mkdir(parents=True, exist_ok=True)
    destination = images_dir / _safe_name(input_path.name, f"input{input_path.suffix.lower()}")
    shutil.copy2(input_path, destination)
    return "", [{"kind": "input-image", "path": destination.relative_to(output_dir).as_posix()}], []


def _annotate_visuals(output_dir: Path, visuals: list[dict[str, Any]], warnings: list[str]) -> None:
    """Attach local dimensions and conservative readability hints."""
    try:
        from PIL import Image
    except ImportError:
        warnings.append("Pillow is unavailable; screenshot dimensions were not inspected.")
        return
    for visual in visuals:
        path = output_dir / str(visual.get("path", ""))
        try:
            with Image.open(path) as image:
                width, height = image.size
            visual["width"] = width
            visual["height"] = height
            visual["review_hint"] = "retake-or-original-resolution" if width < 900 or height < 500 else "inspect-visually"
        except Exception as exc:
            visual["review_hint"] = "unreadable"
            warnings.append(f"Could not inspect screenshot dimensions for {path.name}: {exc}")

def prepare_report(input_path: Path, output_dir: Path, max_pages: int = 80) -> Path:
    """Extract local text and visuals and write a review manifest.

    Args:
        input_path: Existing supported report or image file.
        output_dir: Directory that receives text, images, and manifest JSON.
        max_pages: Maximum PDF pages to render and inspect.

    Returns:
        Path to the generated manifest.

    Raises:
        ValueError: If arguments are invalid.
        ReportPreparationError: If the document cannot be parsed.
    """
    if not input_path.is_file():
        raise ValueError(f"Input file does not exist: {input_path}")
    suffix = input_path.suffix.lower()
    if suffix not in SUPPORTED_SUFFIXES:
        raise ValueError(f"Unsupported file type '{suffix}'. Supported: {', '.join(sorted(SUPPORTED_SUFFIXES))}")
    if max_pages < 1 or max_pages > 500:
        raise ValueError("max_pages must be between 1 and 500")

    output_dir.mkdir(parents=True, exist_ok=True)
    if suffix == ".docx":
        text, visuals, warnings = _prepare_docx(input_path, output_dir)
    elif suffix == ".pdf":
        text, visuals, warnings = _prepare_pdf(input_path, output_dir, max_pages)
    elif suffix in {".txt", ".md"}:
        text, visuals, warnings = _prepare_text(input_path)
    else:
        text, visuals, warnings = _prepare_image(input_path, output_dir)

    _annotate_visuals(output_dir, visuals, warnings)
    contact_sheet = _create_contact_sheet(output_dir, visuals, warnings)
    try:
        from domain_router import route_domain
        domain_routing = route_domain(text)
    except Exception as exc:
        domain_routing = {"selected": None, "confidence": "low", "reason": f"Domain routing unavailable: {exc}"}
        warnings.append(domain_routing["reason"])
    text_path = _write_text(text, output_dir)
    manifest = {
        "skill": SKILL_NAME,
        "version": VERSION,
        "local_only": True,
        "input": {
            "name": input_path.name,
            "suffix": suffix,
            "size_bytes": input_path.stat().st_size,
            "sha256": _sha256(input_path),
        },
        "text_path": text_path.relative_to(output_dir).as_posix(),
        "text_characters": len(text),
        "visual_count": len(visuals),
        "visuals": visuals,
        "visual_overview": contact_sheet,
        "domain_routing": domain_routing,
        "review_signals": classify_report_signals(text, len(visuals)),
        "warnings": warnings,
        "next_step": "The agent must read the text and visually inspect relevant images before classifying or reviewing the report.",
    }
    manifest_path = output_dir / "manifest.json"
    manifest_path.write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
    return manifest_path


def prerequisite_report() -> dict[str, Any]:
    """Return local runtime readiness without making network requests."""
    try:
        import fitz

        fitz_version = getattr(fitz, "version", ("available",))[0]
        pdf_ok = True
    except ImportError:
        fitz_version = "not installed"
        pdf_ok = False
    python_ok = sys.version_info >= (3, 10)
    return {
        "ready": python_ok,
        "checks": [
            {"check": "python", "required": ">=3.10", "found": sys.version.split()[0], "ok": python_ok},
            {"check": "PyMuPDF", "required": "PDF only", "found": str(fitz_version), "ok": pdf_ok},
            {"check": "network", "required": "not required", "found": "disabled by design", "ok": True},
            {"check": "model API key", "required": "not required", "found": "not used", "ok": True},
        ],
    }


def diagnostics() -> dict[str, Any]:
    """Return machine-readable harness metadata."""
    return {
        "skill": SKILL_NAME,
        "version": VERSION,
        "harness_level": "local-document-preparation",
        "commands": ["inspect", "check-prereqs", "diagnostics"],
        "harness_features": {
            "input_validation": True,
            "structured_errors": True,
            "local_only": True,
            "docx_text_and_images": True,
            "pdf_text_images_and_page_renders": True,
        },
    }


def main() -> int:
    """Parse CLI arguments and prepare one report for agent review."""
    parser = argparse.ArgumentParser(description="Prepare a university experiment report for local agent review.")
    parser.add_argument("--input", type=Path, help="Path to a DOCX, PDF, text, Markdown, or image file.")
    parser.add_argument("--output-dir", type=Path, help="Directory for manifest, extracted text, and images.")
    parser.add_argument("--max-pages", type=int, default=80, help="Maximum PDF pages to render (1-500).")
    parser.add_argument("--check-prereqs", action="store_true", help="Report local prerequisites as JSON.")
    parser.add_argument("--diagnostics", action="store_true", help="Report harness metadata as JSON.")
    args = parser.parse_args()

    if args.check_prereqs:
        print(json.dumps(prerequisite_report(), ensure_ascii=False, indent=2))
        return 0
    if args.diagnostics:
        print(json.dumps(diagnostics(), ensure_ascii=False, indent=2))
        return 0
    if args.input is None or args.output_dir is None:
        _structured_error(
            "--input and --output-dir are required for report preparation.",
            "validation",
            "Provide both paths, or use --check-prereqs/--diagnostics.",
        )
        return 1

    try:
        manifest_path = prepare_report(args.input.resolve(), args.output_dir.resolve(), args.max_pages)
    except ValueError as exc:
        _structured_error(str(exc), "validation", "Check the file path, extension, and numeric limits.")
        return 1
    except ReportPreparationError as exc:
        _structured_error(str(exc), "runtime", "Open the original file visually or convert it to DOCX/PDF locally.")
        return 1
    except OSError as exc:
        _structured_error(str(exc), "runtime", "Check local file permissions and available disk space.")
        return 1

    print(json.dumps({"ok": True, "manifest": str(manifest_path)}, ensure_ascii=False))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
