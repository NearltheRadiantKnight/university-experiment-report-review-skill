#!/usr/bin/env python3
"""Build structured, style-safe experiment report deliverables from a local DOCX."""
from __future__ import annotations
import argparse, hashlib, json, re, shutil, sys
from collections import Counter
from datetime import datetime
from pathlib import Path
from typing import Any
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.text.paragraph import Paragraph
from validate_plan import PlanValidationError, load_and_validate

SKILL_NAME="university-experiment-report-review-skill"
VERSION="1.5.0"
DEFAULT_FONT="Microsoft YaHei"
CATEGORY_STYLES={
 "guidance":("2F75B5","执行指导"),"evidence":("008C95","证据要求"),
 "writing":("2F75B5","写作改进"),
 "warning":("C65911","注意事项"),"issue":("C00000","发现问题"),
 "suggestion":("1F4E79","修改建议"),
 "example":("548235","参考写法"),"praise":("8064A2","保留优点"),
 "summary":("44546A","总结")}
PRIORITY_LABELS={"blocker":"阻塞","high":"高","medium":"中","low":"低","optional":"可选"}
PRIORITY_ORDER={"blocker":0,"high":1,"medium":2,"low":3,"optional":4}
DEFAULT_MINUTES={"blocker":15,"high":15,"medium":10,"low":5,"optional":5}
BUDGET_LIMITS={"15m":15,"1h":60,"half_day":240,"full":None}
SIGNAL_LABELS={"green":"绿灯 · 可以提交","yellow":"黄灯 · 小修后提交","red":"红灯 · 暂不建议提交"}
ISSUE_CATEGORIES={"issue","warning","suggestion","guidance","evidence","writing"}
STRENGTH_CATEGORIES={"praise","example"}

class ReportBuildError(RuntimeError): pass

def _error_payload(message:str,error_type:str,hint:str)->None:
 print(json.dumps({"error":message,"error_type":error_type,"hint":hint},ensure_ascii=False),file=sys.stderr)

def _sha256(path:Path)->str:
 digest=hashlib.sha256()
 with path.open("rb") as stream:
  for block in iter(lambda:stream.read(1024*1024),b""): digest.update(block)
 return digest.hexdigest()

def _safe_stem(text:str)->str:
 return re.sub(r'[\\/:*?"<>|]+',"-",text).strip(" .-") or "experiment-report"

def _unique_path(path:Path,job_id:str)->Path:
 return path if not path.exists() else path.with_name(f"{path.stem}-{job_id}{path.suffix}")

def _all_paragraphs(document:Document)->list[Paragraph]:
 paragraphs=list(document.paragraphs)
 for table in document.tables:
  for row in table.rows:
   for cell in row.cells: paragraphs.extend(cell.paragraphs)
 return paragraphs

def _ensure_styles(document:Document)->None:
 styles=document.styles
 if "Review Added" not in styles:
  style=styles.add_style("Review Added",WD_STYLE_TYPE.PARAGRAPH)
  style.base_style=styles["Normal"]
  style.paragraph_format.space_before=Pt(4); style.paragraph_format.space_after=Pt(4)
 if "Review Item" not in styles:
  style=styles.add_style("Review Item",WD_STYLE_TYPE.PARAGRAPH)
  style.base_style=styles["Review Added"]
  style.paragraph_format.left_indent=Pt(18); style.paragraph_format.first_line_indent=Pt(-12)

def _insert_after(paragraph:Paragraph,style_name:str="Review Added")->Paragraph:
 element=OxmlElement("w:p"); paragraph._p.addnext(element); inserted=Paragraph(element,paragraph._parent)
 try: inserted.style=style_name
 except KeyError: pass
 return inserted

def _set_run_font(run:Any,font_name:str,color_hex:str,size_pt:float,bold:bool=False)->None:
 run.font.name=font_name; run.font.color.rgb=RGBColor.from_string(color_hex); run.font.size=Pt(size_pt); run.bold=bold
 rpr=run._element.get_or_add_rPr(); fonts=rpr.rFonts
 if fonts is None: fonts=OxmlElement("w:rFonts"); rpr.insert(0,fonts)
 for key in ("w:ascii","w:hAnsi","w:eastAsia"): fonts.set(qn(key),font_name)

def _add_runs(paragraph:Paragraph,addition:dict[str,Any],label_override:str|None=None,text_override:str|None=None)->None:
 category=str(addition.get("category","suggestion")); color,default_label=CATEGORY_STYLES.get(category,CATEGORY_STYLES["suggestion"])
 label=label_override or str(addition.get("label",default_label)).strip() or default_label
 text=str(addition.get("text","") if text_override is None else text_override).strip()
 font=str(addition.get("font_name",DEFAULT_FONT)).strip() or DEFAULT_FONT; size=float(addition.get("font_size_pt",10.5))
 label_run=paragraph.add_run(f"【AI 审阅·{label}】"); _set_run_font(label_run,font,color,size,True)
 if text:
  body=paragraph.add_run(text); _set_run_font(body,font,color,size)

def _apply_table_style(table: Any) -> None:
 try:
  table.style="Table Grid"
 except KeyError:
  pass
def _item_text(addition:dict[str,Any],item:Any,checklist:bool)->str:
 return f"{'☐' if checklist else '•'} {str(item).strip()}"

def _render_after(document:Document,anchor:Paragraph,addition:dict[str,Any])->Paragraph:
 block=str(addition.get("block_type","paragraph")); current=_insert_after(anchor); _add_runs(current,addition)
 if block in {"bullets","checklist"}:
  for item in addition.get("items",[]):
   current=_insert_after(current,"Review Item")
   category=str(addition.get("category","suggestion")); color=CATEGORY_STYLES.get(category,CATEGORY_STYLES["suggestion"])[0]
   run=current.add_run(_item_text(addition,item,block=="checklist")); _set_run_font(run,str(addition.get("font_name",DEFAULT_FONT)),color,float(addition.get("font_size_pt",10.5)))
 elif block=="table":
  tail=_insert_after(current); table=document.add_table(rows=1,cols=len(addition["columns"])); _apply_table_style(table)
  for index,value in enumerate(addition["columns"]):
   cell=table.rows[0].cells[index]; cell.text=str(value)
   for run in cell.paragraphs[0].runs: _set_run_font(run,DEFAULT_FONT,"44546A",9,True)
  for values in addition.get("rows",[]):
   cells=table.add_row().cells
   for index,value in enumerate(values):
    cells[index].text="" if value is None else str(value)
    for run in cells[index].paragraphs[0].runs: _set_run_font(run,DEFAULT_FONT,"1F1F1F",9)
  tail._p.addprevious(table._tbl); current=tail
 return current

def _render_append(document:Document,addition:dict[str,Any])->None:
 block=str(addition.get("block_type","paragraph")); paragraph=document.add_paragraph(style="Review Added"); _add_runs(paragraph,addition)
 if block in {"bullets","checklist"}:
  for item in addition.get("items",[]):
   row=document.add_paragraph(style="Review Item"); color=CATEGORY_STYLES.get(str(addition.get("category")),CATEGORY_STYLES["suggestion"])[0]
   run=row.add_run(_item_text(addition,item,block=="checklist")); _set_run_font(run,DEFAULT_FONT,color,float(addition.get("font_size_pt",10.5)))
 elif block=="table":
  table=document.add_table(rows=1,cols=len(addition["columns"])); _apply_table_style(table)
  for index,value in enumerate(addition["columns"]): table.rows[0].cells[index].text=str(value)
  for values in addition.get("rows",[]):
   cells=table.add_row().cells
   for index,value in enumerate(values): cells[index].text="" if value is None else str(value)

def _find_anchor(paragraphs:list[Paragraph],addition:dict[str,Any])->Paragraph|None:
 if "paragraph_index" in addition:
  index=int(addition["paragraph_index"])
  if index<0 or index>=len(paragraphs): raise ReportBuildError(f"paragraph_index {index} is outside the source document.")
  return paragraphs[index]
 text=str(addition.get("anchor_text","")).strip(); occurrence=int(addition.get("occurrence",1))
 if not text: return None
 matches=[paragraph for paragraph in paragraphs if text in paragraph.text]
 return matches[occurrence-1] if len(matches)>=occurrence else None

def _append_summary(document:Document,plan:dict[str,Any])->None:
 if not plan.get("include_summary_appendix"): return
 paragraph=document.add_paragraph(); paragraph.add_run().add_break(WD_BREAK.PAGE)
 _render_append(document,{"category":"summary","label":"提交判断","text":str(plan.get("verdict",""))})
 _render_append(document,{"category":"summary","label":"整体说明","text":str(plan.get("summary",""))})

def _addition_plain_text(addition:dict[str,Any])->str:
 block=str(addition.get("block_type","paragraph")); text=str(addition.get("text","")).strip()
 if block in {"bullets","checklist"}: return "；".join(([text] if text else [])+[str(x) for x in addition.get("items",[])])
 if block=="table": return f"表格：{', '.join(str(x) for x in addition.get('columns',[]))}（{len(addition.get('rows',[]))} 行）"
 return text

def _estimated_minutes(addition:dict[str,Any])->int:
 return int(addition.get("estimated_minutes",DEFAULT_MINUTES.get(str(addition.get("priority","medium")),10)))

def _select_budgeted_additions(plan:dict[str,Any])->list[dict[str,Any]]:
 additions=[dict(item) for item in plan.get("additions",[])]
 for item in additions: item["estimated_minutes"]=_estimated_minutes(item)
 limit=BUDGET_LIMITS.get(str(plan.get("time_budget","full")))
 if limit is None: return additions
 ranked=sorted(enumerate(additions),key=lambda pair:(PRIORITY_ORDER.get(str(pair[1].get("priority","medium")),2),pair[0]))
 selected=[]; used=0
 for index,item in ranked:
  minutes=item["estimated_minutes"]
  if len(selected)>=5: break
  if used+minutes<=limit: selected.append((index,item)); used+=minutes
 return [item for _,item in sorted(selected,key=lambda pair:pair[0])]

def _add_table(document:Document,headers:list[str],rows:list[list[Any]])->None:
 table=document.add_table(rows=1,cols=len(headers)); _apply_table_style(table)
 for cell,value in zip(table.rows[0].cells,headers):
  cell.text=str(value)
  for run in cell.paragraphs[0].runs: _set_run_font(run,DEFAULT_FONT,"44546A",9,True)
 for values in rows:
  cells=table.add_row().cells
  for index,value in enumerate(values):
   cells[index].text="" if value is None else str(value)
   for run in cells[index].paragraphs[0].runs: _set_run_font(run,DEFAULT_FONT,"1F1F1F",9)

def _append_review_appendix(document:Document,plan:dict[str,Any],actions:list[dict[str,Any]],output_dir:Path,job_id:str,plan_dir:Path)->list[dict[str,Any]]:
 page=document.add_paragraph(); page.add_run().add_break(WD_BREAK.PAGE)
 document.add_heading(str(plan.get("appendix_title","AI 综合审阅附录")),level=1)
 signal=str(plan.get("submission_signal") or ("red" if plan.get("source_state")=="blank" else "yellow")); time_budget=str(plan.get("time_budget","full"))
 _render_append(document,{"category":"summary","label":"提交状态","text":SIGNAL_LABELS.get(signal,signal)})
 _render_append(document,{"category":"summary","label":"时间预算","text":f"{time_budget} · 预计 {sum(_estimated_minutes(item) for item in actions)} 分钟"})
 _render_append(document,{"category":"summary","label":"提交判断","text":str(plan.get("verdict",""))})
 _render_append(document,{"category":"summary","label":"整体说明","text":str(plan.get("summary",""))})
 if actions:
  document.add_heading("预算内优先行动",level=2)
  rows=[]
  for item in actions[:5]:
   category=str(item.get("category","suggestion")); default=CATEGORY_STYLES.get(category,CATEGORY_STYLES["suggestion"])[1]
   rows.append([PRIORITY_LABELS.get(str(item.get("priority","medium")),"中"),f"{_estimated_minutes(item)} 分钟",str(item.get("label",default)),_addition_plain_text(item)])
  _add_table(document,["优先级","预计耗时","行动","完成标准"],rows)
 for field,title in (("false_completion_findings","伪完成识别"),("contamination_findings","污染内容检查")):
  findings=plan.get(field,[])
  if findings:
   document.add_heading(title,level=2)
   _add_table(document,["严重度","位置","问题","证据","最小修复动作"],[[PRIORITY_LABELS.get(str(item.get("severity","medium")),"中"),item.get("location",""),item.get("issue",""),item.get("evidence",""),item.get("action","")] for item in findings])
 screenshots=[]; asset_dir=output_dir/"review-assets"/job_id
 for index,item in enumerate(plan.get("screenshot_evidence",[]),1):
  entry=dict(item); image_path=str(item.get("image_path","")).strip()
  if image_path:
   candidate=Path(image_path).expanduser(); candidate=candidate if candidate.is_absolute() else plan_dir/candidate
   if candidate.is_file() and candidate.suffix.lower() in {".png",".jpg",".jpeg",".bmp",".gif",".tif",".tiff"}:
    asset_dir.mkdir(parents=True,exist_ok=True); destination=asset_dir/f"screenshot-{index:03d}{candidate.suffix.lower()}"; shutil.copy2(candidate,destination)
    entry["preview_name"]=destination.relative_to(output_dir).as_posix()
  screenshots.append(entry)
 if screenshots:
  document.add_heading("截图证据与重拍指令",level=2)
  _add_table(document,["截图","证明目标","直接可见","可读性","正文匹配","重拍指令"],[[item.get("source_ref",""),item.get("claim",""),item.get("observed",""),item.get("readability",""),item.get("matches_text",""),item.get("retake_instruction","")] for item in screenshots])
  for item in screenshots:
   if item.get("embed_in_appendix") and item.get("preview_name"):
    document.add_paragraph(str(item.get("source_ref","独立截图")))
    document.add_picture(str(output_dir/item["preview_name"]),width=Inches(5.8))
 return screenshots
def build_report(source:Path,plan_path:Path,output_dir:Path)->tuple[Path,Path]:
 if not source.is_file(): raise ValueError(f"Source file does not exist: {source}")
 if source.suffix.lower()!=".docx": raise ValueError("Style-preserving generated reports require a .docx source file.")
 try: plan,warnings=load_and_validate(plan_path)
 except PlanValidationError as exc: raise ReportBuildError(str(exc)) from exc
 if plan.get("output_mode","single_docx")=="guidance_only": raise ReportBuildError("output_mode=guidance_only does not generate a DOCX; return the reviewed plan to the user instead.")
 try: document=Document(str(source))
 except Exception as exc: raise ReportBuildError(f"DOCX could not be opened: {exc}") from exc
 output_dir.mkdir(parents=True,exist_ok=True); job_id=datetime.now().strftime("%Y%m%d-%H%M%S-%f")
 selected_additions=_select_budgeted_additions(plan)
 _ensure_styles(document); original_paragraphs=_all_paragraphs(document); last_by_anchor={}; appendix=[]; applied=[]
 strict=bool(plan.get("strict_anchors",True))
 for index,addition in enumerate(selected_additions,1):
  position=str(addition.get("position","after")); anchor=_find_anchor(original_paragraphs,addition)
  if position=="append": appendix.append(addition); applied.append({"index":index,"location":"appendix","anchor_found":True}); continue
  if anchor is None:
   if strict or str(addition.get("on_anchor_missing","error"))=="error": raise ReportBuildError(f"Addition {index} anchor was not found. Fix the plan instead of creating an unlocated section.")
   appendix.append(addition); applied.append({"index":index,"location":"appendix-fallback","anchor_found":False}); continue
  key=id(anchor._p); insertion=last_by_anchor.get(key,anchor); last_by_anchor[key]=_render_after(document,insertion,addition)
  applied.append({"index":index,"location":"after-anchor","anchor_found":True,"anchor_text":anchor.text[:120]})
 screenshots=_append_review_appendix(document,plan,selected_additions,output_dir,job_id,plan_path.parent)
 if appendix:
  document.add_heading("定位到附录的修改",level=2)
  for addition in appendix: _render_append(document,addition)
 report_label="实验执行报告" if plan["report_kind"]=="execution" else "修改报告"
 output_path=_unique_path(output_dir/f"{_safe_stem(source.stem)}-{report_label}.docx",job_id); document.save(output_path)
 generated_files=[{"kind":"annotated","label":report_label,"name":output_path.name,"sha256":_sha256(output_path)}]
 issue_additions=[item for item in selected_additions if str(item.get("category","suggestion")) in ISSUE_CATEGORIES]
 strength_additions=[item for item in selected_additions if str(item.get("category","suggestion")) in STRENGTH_CATEGORIES]
 counts=Counter(str(item.get("priority","medium")) for item in issue_additions); actions=[]; strengths=[]
 for item in issue_additions[:5]:
  category=str(item.get("category","suggestion")); priority=str(item.get("priority","medium")); _hex,default_label=CATEGORY_STYLES.get(category,CATEGORY_STYLES["suggestion"])
  actions.append({"label":str(item.get("label",default_label)),"priority":priority,"priority_label":PRIORITY_LABELS.get(priority,"中"),"category":category,"evidence_basis":str(item.get("evidence_basis","unverified")),"estimated_minutes":_estimated_minutes(item),"text":_addition_plain_text(item)[:500]})
 for item in strength_additions[:5]:
  category=str(item.get("category","praise")); _hex,default_label=CATEGORY_STYLES.get(category,CATEGORY_STYLES["praise"])
  strengths.append({"label":str(item.get("label",default_label)),"category":category,"text":_addition_plain_text(item)[:500]})
 metadata={"job_id":job_id,"skill":SKILL_NAME,"version":VERSION,"created_at":datetime.now().astimezone().isoformat(timespec="seconds"),"report_kind":plan["report_kind"],"report_label":report_label,"source_state":plan["source_state"],"source_name":source.name,"domain_profile":plan.get("domain_profile"),"domain_confidence":plan.get("domain_confidence"),"domain_profile_basis":plan.get("domain_profile_basis"),"source_sha256":_sha256(source),"generated_name":output_path.name,"generated_sha256":_sha256(output_path),"generated_files":generated_files,"summary":str(plan.get("summary","")).strip(),"verdict":str(plan.get("verdict","")).strip(),"submission_signal":plan.get("submission_signal") or ("red" if plan.get("source_state")=="blank" else "yellow"),"time_budget":plan.get("time_budget","full"),"review_depth":plan.get("review_depth","standard"),"review_focus":plan.get("review_focus","comprehensive"),"output_mode":plan.get("output_mode","single_docx"),"estimated_minutes":sum(_estimated_minutes(item) for item in issue_additions),"addition_count":len(selected_additions),"priority_counts":dict(counts),"actions":actions,"strengths":strengths,"false_completion_findings":plan.get("false_completion_findings",[]),"contamination_findings":plan.get("contamination_findings",[]),"screenshot_evidence":screenshots,"plan_warnings":warnings,"applied":applied,"style_policy":{"original_content":"preserved","generated_content":"structured, labeled, and category-colored"},"local_only":True}
 metadata_path=output_dir/f"{job_id}.metadata.json"; metadata_path.write_text(json.dumps(metadata,ensure_ascii=False,indent=2),encoding="utf-8")
 return output_path,metadata_path
def diagnostics()->dict[str,Any]:
 return {"skill":SKILL_NAME,"version":VERSION,"commands":["build-report","diagnostics"],"features":{"structured_blocks":True,"strict_anchors":True,"numeric_score_gate":True,"single_docx_delivery":True,"external_model_api":False}}

def main()->int:
 parser=argparse.ArgumentParser(description="Build structured experiment-report deliverables."); parser.add_argument("--source",type=Path); parser.add_argument("--plan",type=Path); parser.add_argument("--output-dir",type=Path); parser.add_argument("--diagnostics",action="store_true"); args=parser.parse_args()
 if args.diagnostics: print(json.dumps(diagnostics(),ensure_ascii=False,indent=2)); return 0
 if None in (args.source,args.plan,args.output_dir): _error_payload("--source, --plan, and --output-dir are required.","validation","Supply all paths."); return 1
 try: output,metadata=build_report(args.source.resolve(),args.plan.resolve(),args.output_dir.resolve())
 except (ValueError,ReportBuildError,OSError) as exc: _error_payload(str(exc),"runtime","Check paths, plan quality gates, anchors, and open files."); return 1
 print(json.dumps({"ok":True,"output":str(output),"metadata":str(metadata)},ensure_ascii=False)); return 0
if __name__=="__main__": raise SystemExit(main())
