#!/usr/bin/env python3
"""Run structural QA and optional cross-platform render QA for generated DOCX files."""
from __future__ import annotations
import json, os, shutil, subprocess, sys, zipfile
from pathlib import Path
from typing import Any
from docx import Document

class ReportQualityError(RuntimeError): pass

def _media_count(path:Path)->int:
 with zipfile.ZipFile(path) as archive: return sum(1 for name in archive.namelist() if name.startswith("word/media/") and not name.endswith("/"))

def _run(command:list[str],timeout:int=120)->subprocess.CompletedProcess[str]:
 return subprocess.run(command,capture_output=True,text=True,timeout=timeout,check=False)

def _render_external(docx:Path,render_dir:Path,pdf_path:Path)->tuple[bool,str]:
 renderer=os.environ.get("EXPERIMENT_REPORT_RENDERER","").strip()
 if not renderer: return False,"EXPERIMENT_REPORT_RENDERER is not set."
 path=Path(renderer).expanduser()
 if not path.is_file(): return False,f"Configured renderer does not exist: {path}"
 result=_run([sys.executable,str(path),str(docx),"--output_dir",str(render_dir),"--emit_pdf"])
 candidates=list(render_dir.glob("*.pdf"))
 if result.returncode==0 and candidates:
  shutil.copy2(candidates[0],pdf_path); return True,(result.stdout+result.stderr).strip()
 return False,(result.stdout+result.stderr).strip() or f"Renderer exited {result.returncode}."

def _render_soffice(docx:Path,render_dir:Path,pdf_path:Path)->tuple[bool,str]:
 binary=shutil.which("soffice") or shutil.which("libreoffice")
 if not binary: return False,"LibreOffice/soffice is not installed."
 result=_run([binary,"--headless","--convert-to","pdf","--outdir",str(render_dir),str(docx)])
 candidate=render_dir/f"{docx.stem}.pdf"
 if result.returncode==0 and candidate.is_file(): shutil.copy2(candidate,pdf_path); return True,(result.stdout+result.stderr).strip()
 return False,(result.stdout+result.stderr).strip() or f"LibreOffice exited {result.returncode}."

def _render_word_com(docx:Path,pdf_path:Path)->tuple[bool,str]:
 if os.name!="nt" or not shutil.which("powershell"): return False,"Windows Word COM is unavailable on this platform."
 source=str(docx.resolve()).replace("'","''"); target=str(pdf_path.resolve()).replace("'","''")
 script=("$ErrorActionPreference='Stop'; $word=$null; $doc=$null; try { "
  "$word=New-Object -ComObject Word.Application; $word.Visible=$false; $word.DisplayAlerts=0; "
  f"$doc=$word.Documents.Open('{source}', $false, $true); $doc.ExportAsFixedFormat('{target}',17); "
  "} finally { if($doc){$doc.Close($false)}; if($word){$word.Quit()} }")
 result=_run(["powershell","-NoProfile","-STA","-ExecutionPolicy","Bypass","-Command",script])
 return result.returncode==0 and pdf_path.is_file(),(result.stdout+result.stderr).strip()

def _word_com_enabled()->bool:
 return os.environ.get("EXPERIMENT_REPORT_ALLOW_WORD_COM","").strip().lower() in {"1","true","yes","on"}

def _rasterize(pdf_path:Path,page_dir:Path)->list[Path]:
 try: import fitz
 except ImportError as exc: raise ReportQualityError("PyMuPDF is required to rasterize rendered PDFs.") from exc
 page_dir.mkdir(parents=True,exist_ok=True); document=fitz.open(pdf_path); pages=[]
 try:
  for index in range(len(document)):
   path=page_dir/f"page-{index+1:03d}.png"; document.load_page(index).get_pixmap(matrix=fitz.Matrix(1.5,1.5),alpha=False).save(path); pages.append(path)
 finally: document.close()
 return pages

def _contact_sheet(pages:list[Path],output:Path)->Path|None:
 if not pages: return None
 try: from PIL import Image,ImageDraw,ImageFont
 except ImportError: return None
 width,height,columns=560,760,2; rows=(len(pages)+1)//2; sheet=Image.new("RGB",(width*columns,height*rows),"white"); draw=ImageDraw.Draw(sheet); font=ImageFont.load_default()
 for index,path in enumerate(pages):
  with Image.open(path) as image:
   image=image.convert("RGB"); image.thumbnail((width-30,height-45)); x=(index%2)*width+(width-image.width)//2; y=(index//2)*height+28; sheet.paste(image,(x,y))
  draw.text(((index%2)*width+12,(index//2)*height+8),f"Page {index+1}",fill="#17324d",font=font)
 sheet.save(output,quality=88); return output

def render_report(docx:Path,job_id:str)->dict[str,Any]:
 render_dir=docx.parent/"render-work"/job_id; render_dir.mkdir(parents=True,exist_ok=True); pdf_path=docx.parent/f"{job_id}.render.pdf"
 attempts=[]; backend=None
 renderers=[("external",lambda:_render_external(docx,render_dir,pdf_path)),("libreoffice",lambda:_render_soffice(docx,render_dir,pdf_path))]
 if _word_com_enabled(): renderers.append(("word-com",lambda:_render_word_com(docx,pdf_path)))
 for name,function in renderers:
  try: ok,message=function()
  except (OSError,subprocess.SubprocessError) as exc: ok,message=False,str(exc)
  attempts.append({"backend":name,"ok":ok,"message":message[-1200:]})
  if ok: backend=name; break
 if not _word_com_enabled(): attempts.append({"backend":"word-com","ok":False,"message":"Skipped by default to avoid Windows sandbox/COM popups. Opt in with --allow-word-com."})
 if not backend:
  permission_required=any("80070520" in item["message"] or "logon session" in item["message"].lower() for item in attempts)
  status="permission-required" if permission_required else "unavailable"
  resolution="Microsoft Word 已安装，但当前沙箱会话不能访问桌面 COM。请在允许桌面权限的主机会话中重跑流水线。" if permission_required else "安装 LibreOffice，或设置 EXPERIMENT_REPORT_RENDERER 指向可用的本地 DOCX 渲染器。"
  return {"status":status,"backend":None,"attempts":attempts,"pdf":None,"pages":[],"preview":None,"resolution":resolution}
 pages=_rasterize(pdf_path,render_dir/"pages"); preview=_contact_sheet(pages,docx.parent/f"{job_id}.render-preview.jpg")
 return {"status":"passed","backend":backend,"attempts":attempts,"pdf":pdf_path.name,"pages":[str(path) for path in pages],"page_count":len(pages),"preview":preview.name if preview else None}

def check_report(source:Path,generated:Path,metadata_path:Path,require_render:bool=False)->tuple[dict[str,Any],Path]:
 metadata=json.loads(metadata_path.read_text(encoding="utf-8")); document=Document(generated); text="\n".join(p.text for p in document.paragraphs); expected=int(metadata.get("addition_count",0)); labels=text.count("【Codex 新增·")
 checks={"docx_opens":True,"source_media_preserved":_media_count(generated)>=_media_count(source),"all_additions_labeled":labels>=expected,"no_unlocated_marker":"未定位内容" not in text,"paragraph_length_ok":all(len(p.text)<=700 for p in document.paragraphs)}
 structural_ok=all(checks.values())
 if "自动质量检查" not in text:
  document.add_heading("自动质量检查",level=2); table=document.add_table(rows=1,cols=2)
  try: table.style="Table Grid"
  except KeyError: pass
  table.rows[0].cells[0].text="检查项"; table.rows[0].cells[1].text="结果"
  labels_map={"docx_opens":"DOCX 可打开","source_media_preserved":"原图片保留","all_additions_labeled":"新增内容有标签","no_unlocated_marker":"无未定位内容","paragraph_length_ok":"段落长度合理"}
  for key,value in checks.items(): cells=table.add_row().cells; cells[0].text=labels_map[key]; cells[1].text="通过" if value else "未通过"
  document.save(generated)
 render=render_report(generated,str(metadata["job_id"])); visual_ok=render["status"]=="passed"; shipping_ready=structural_ok and (visual_ok or not require_render)
 payload={"ok":structural_ok,"shipping_ready":shipping_ready,"render_required":require_render,"checks":checks,"render":render,"generated":generated.name}
 qa_path=generated.parent/f"{metadata['job_id']}.quality.json"; qa_path.write_text(json.dumps(payload,ensure_ascii=False,indent=2),encoding="utf-8")
 if not structural_ok: raise ReportQualityError(f"Generated report failed structural checks: {checks}")
 if require_render and not visual_ok: raise ReportQualityError("Render QA is required but no renderer completed successfully.")
 return payload,qa_path
